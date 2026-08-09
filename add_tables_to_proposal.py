#!/usr/bin/env python3
"""
Add Professional Tables to BORS Proposal
Creates tables for Business Model Canvas, Financial data, etc.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_business_model_canvas_table(doc):
    """Create Business Model Canvas as a table"""
    
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == 'BUSINESS MODEL CANVAS':
            # Find where to insert table (after the intro paragraph)
            insert_pos = i + 3
            
            # Create table
            table = doc.add_table(rows=9, cols=2)
            # table.style = 'Light Grid Accent 1'  # Skip style, will format manually
            table.autofit = False
            table.allow_autofit = False
            
            # Set column widths
            for row in table.rows:
                row.cells[0].width = Inches(2.5)
                row.cells[1].width = Inches(4.0)
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Component'
            header_cells[1].text = 'Description'
            
            # Style header
            for cell in header_cells:
                set_cell_background(cell, '2E7D32')  # Green
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(12)
            
            # Data rows
            data = [
                ('Value Proposition', 
                 '• Efficient digital record management\n'
                 '• 70% faster certificate processing\n'
                 '• Real-time communication platform\n'
                 '• 24/7 accessible services\n'
                 '• Secure data storage with encryption'),
                
                ('Customer Segments',
                 'Primary: Barangay Officials (8 roles), Residents\n'
                 'Secondary: Municipal Government, Partner Organizations'),
                
                ('Channels',
                 '• Web Browser (Desktop & Mobile)\n'
                 '• Direct URL: https://barangay.mocogo.site\n'
                 '• Email notifications\n'
                 '• Barangay hall kiosks'),
                
                ('Customer Relationships',
                 '• Self-service portal\n'
                 '• Automated notifications\n'
                 '• Real-time status tracking\n'
                 '• 24/7 technical support'),
                
                ('Key Partners',
                 '• Municipal Government\n'
                 '• Internet Service Providers\n'
                 '• Web Hosting Companies\n'
                 '• IT Support Contractors'),
                
                ('Key Resources',
                 'Human: Developers, Designers, Support Staff\n'
                 'Technical: Servers, Database, Domain, SSL\n'
                 'Knowledge: Documentation, Training Materials'),
                
                ('Key Activities',
                 '• System development and maintenance\n'
                 '• User training and support\n'
                 '• Data backup and security\n'
                 '• Feature enhancements'),
                
                ('Revenue Streams',
                 'Initial: Barangay budget (PHP 401,150)\n'
                 'Annual: Maintenance budget (PHP 106,500)\n'
                 'Savings: PHP 100,000/year cost reduction'),
            ]
            
            for idx, (component, description) in enumerate(data, start=1):
                row = table.rows[idx]
                row.cells[0].text = component
                row.cells[1].text = description
                
                # Style component cell
                set_cell_background(row.cells[0], 'E8F5E9')  # Light green
                for paragraph in row.cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)
                
                # Style description cell
                for paragraph in row.cells[1].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
            
            # Move table to correct position
            table._element.getparent().remove(table._element)
            doc.paragraphs[insert_pos]._element.addprevious(table._element)
            
            print("✓ Business Model Canvas table created")
            break

def create_financial_cost_table(doc):
    """Create Projected Cost table"""
    
    for i, para in enumerate(doc.paragraphs):
        if 'Projected Cost' in para.text and 'Sample' not in para.text:
            # Create table after this heading
            table = doc.add_table(rows=18, cols=2)
            # table.style = 'Light Grid Accent 1'  # Skip style
            
            # Set column widths
            for row in table.rows:
                row.cells[0].width = Inches(4.0)
                row.cells[1].width = Inches(2.5)
            
            # Header
            header = table.rows[0].cells
            header[0].text = 'Cost Item'
            header[1].text = 'Amount (PHP)'
            
            for cell in header:
                set_cell_background(cell, '1565C0')  # Blue
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Data
            costs = [
                ('DEVELOPMENT COSTS', ''),
                ('System Development', '150,000'),
                ('Database Design & Setup', '20,000'),
                ('UI/UX Design', '30,000'),
                ('Testing & QA', '25,000'),
                ('Documentation', '15,000'),
                ('Subtotal', '240,000'),
                ('', ''),
                ('INFRASTRUCTURE COSTS (Year 1)', ''),
                ('Domain Name & SSL Certificate', '6,500'),
                ('Web Hosting (VPS)', '15,000'),
                ('Database Hosting', '8,000'),
                ('Email Service', '3,000'),
                ('Backup Storage', '4,000'),
                ('Subtotal', '36,500'),
                ('', ''),
                ('TOTAL PROJECT COST', '401,150'),
            ]
            
            for idx, (item, amount) in enumerate(costs, start=1):
                row = table.rows[idx]
                row.cells[0].text = item
                row.cells[1].text = amount
                
                # Style section headers
                if item.isupper() and amount == '':
                    set_cell_background(row.cells[0], 'BBDEFB')
                    for paragraph in row.cells[0].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                    table.rows[idx].cells[0].merge(table.rows[idx].cells[1])
                
                # Style subtotals and total
                elif 'Subtotal' in item or 'TOTAL' in item:
                    for cell in [row.cells[0], row.cells[1]]:
                        set_cell_background(cell, 'E3F2FD')
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                
                # Right-align amounts
                if amount:
                    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Move table
            table._element.getparent().remove(table._element)
            doc.paragraphs[i+1]._element.addprevious(table._element)
            
            print("✓ Financial cost table created")
            break

def create_cash_flow_table(doc):
    """Create 5-Year Cash Flow table"""
    
    for i, para in enumerate(doc.paragraphs):
        if 'Cash Flow' in para.text and 'Sample' not in para.text:
            # Create table
            table = doc.add_table(rows=7, cols=6)
            # table.style = 'Light Grid Accent 1'  # Skip style
            
            # Header
            headers = ['Year', 'Investment', 'Annual Costs', 'Cost Savings', 'Net Cash Flow', 'Cumulative']
            for idx, header_text in enumerate(headers):
                cell = table.rows[0].cells[idx]
                cell.text = header_text
                set_cell_background(cell, '1565C0')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(10)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Data
            cash_flow_data = [
                ('Year 1', '(401,150)', '-', '100,000', '(301,150)', '(301,150)'),
                ('Year 2', '-', '(106,500)', '100,000', '(6,500)', '(307,650)'),
                ('Year 3', '-', '(106,500)', '100,000', '(6,500)', '(314,150)'),
                ('Year 4', '-', '(106,500)', '100,000', '(6,500)', '(320,650)'),
                ('Year 5', '-', '(106,500)', '100,000', '(6,500)', '(327,150)'),
                ('TOTAL', '(401,150)', '(426,000)', '500,000', '(327,150)', '(327,150)'),
            ]
            
            for row_idx, row_data in enumerate(cash_flow_data, start=1):
                for col_idx, value in enumerate(row_data):
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = value
                    
                    # Right-align numbers
                    if col_idx > 0:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    # Style total row
                    if 'TOTAL' in value:
                        set_cell_background(cell, 'E3F2FD')
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    
                    # Font size
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
            
            # Move table
            table._element.getparent().remove(table._element)
            doc.paragraphs[i+1]._element.addprevious(table._element)
            
            print("✓ Cash flow table created")
            break

def create_market_validation_table(doc):
    """Create Market Validation Results table"""
    
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == 'MARKET VALIDATION':
            # Find position after summary
            insert_pos = i + 5
            
            # Create table
            table = doc.add_table(rows=8, cols=3)
            # table.style = 'Light Grid Accent 1'  # Skip style
            
            # Header
            headers = ['Survey Question', 'Response', 'Percentage']
            for idx, header_text in enumerate(headers):
                cell = table.rows[0].cells[idx]
                cell.text = header_text
                set_cell_background(cell, '2E7D32')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Data
            survey_data = [
                ('Willingness to use digital system', 'Yes', '92%'),
                ('Current processing too slow', 'Yes', '78%'),
                ('Want better record management', 'Yes', '85%'),
                ('Have access to devices', 'Yes', '95%'),
                ('Comfortable with web systems', 'Yes', '88%'),
                ('Support digital transformation', 'Yes', '90%'),
                ('Overall satisfaction (current)', 'Satisfied', '35%'),
            ]
            
            for row_idx, (question, response, percentage) in enumerate(survey_data, start=1):
                table.rows[row_idx].cells[0].text = question
                table.rows[row_idx].cells[1].text = response
                table.rows[row_idx].cells[2].text = percentage
                
                # Center align response and percentage
                table.rows[row_idx].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                table.rows[row_idx].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Highlight high percentages
                if percentage and int(percentage.replace('%', '')) >= 85:
                    set_cell_background(table.rows[row_idx].cells[2], 'C8E6C9')
            
            # Move table
            if insert_pos < len(doc.paragraphs):
                table._element.getparent().remove(table._element)
                doc.paragraphs[insert_pos]._element.addprevious(table._element)
            
            print("✓ Market validation table created")
            break

def create_manpower_table(doc):
    """Create Manpower Requirements table"""
    
    for i, para in enumerate(doc.paragraphs):
        if 'Manpower requirements' in para.text:
            # Create table
            table = doc.add_table(rows=11, cols=3)
            # table.style = 'Light Grid Accent 1'  # Skip style
            
            # Header
            headers = ['Phase', 'Position', 'Quantity']
            for idx, header_text in enumerate(headers):
                cell = table.rows[0].cells[idx]
                cell.text = header_text
                set_cell_background(cell, '6A1B9A')  # Purple
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Data
            manpower_data = [
                ('Development', 'Full-Stack Developer', '1'),
                ('Development', 'UI/UX Designer', '1'),
                ('Development', 'Database Administrator', '1'),
                ('Development', 'QA Tester', '1'),
                ('Development', 'Project Manager', '1'),
                ('Implementation', 'System Administrator', '1'),
                ('Implementation', 'Technical Support Staff', '1'),
                ('Implementation', 'Training Coordinator', '1'),
                ('Operational', 'IT Support Personnel', '1 (part-time)'),
                ('Operational', 'Barangay Staff (trained)', 'All staff'),
            ]
            
            for row_idx, (phase, position, qty) in enumerate(manpower_data, start=1):
                table.rows[row_idx].cells[0].text = phase
                table.rows[row_idx].cells[1].text = position
                table.rows[row_idx].cells[2].text = qty
                
                # Center align quantity
                table.rows[row_idx].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Highlight phases
                if row_idx == 1 or (row_idx > 1 and manpower_data[row_idx-1][0] != phase):
                    set_cell_background(table.rows[row_idx].cells[0], 'E1BEE7')
            
            # Move table
            table._element.getparent().remove(table._element)
            doc.paragraphs[i+1]._element.addprevious(table._element)
            
            print("✓ Manpower requirements table created")
            break

print("Adding professional tables to BORS proposal...")
print("Loading document...")

doc = Document('BORS_Technopreneurship_Proposal_FINAL_COMPLETE.docx')

print("\nCreating tables...")
create_business_model_canvas_table(doc)
create_financial_cost_table(doc)
create_cash_flow_table(doc)
create_market_validation_table(doc)
create_manpower_table(doc)

print("\nSaving document with tables...")
doc.save('BORS_Technopreneurship_Proposal_WITH_TABLES.docx')

print("\n✅ Professional tables added!")
print("📄 File: BORS_Technopreneurship_Proposal_WITH_TABLES.docx")
print("\n📊 Tables created:")
print("   ✓ Business Model Canvas (9x2 table)")
print("   ✓ Projected Cost Breakdown (18x2 table)")
print("   ✓ 5-Year Cash Flow (7x6 table)")
print("   ✓ Market Validation Results (8x3 table)")
print("   ✓ Manpower Requirements (11x3 table)")
print("\n🎨 Features:")
print("   • Color-coded headers")
print("   • Professional styling")
print("   • Proper alignment")
print("   • Easy to read")
print("\n✏️  Document is now complete with tables!")
