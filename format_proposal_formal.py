#!/usr/bin/env python3
"""
Format BORS Proposal - Professional Formatting
Fixes spacing, alignment, and makes it formal
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def format_document_professionally(doc):
    """Apply professional formatting to the entire document"""
    
    # Set default paragraph spacing
    for para in doc.paragraphs:
        # Skip empty paragraphs
        if not para.text.strip():
            continue
            
        # Set paragraph format
        para_format = para.paragraph_format
        
        # Line spacing - 1.5 for body text
        para_format.line_spacing = 1.5
        
        # Space after paragraphs
        para_format.space_after = Pt(6)
        
        # First line indent for body paragraphs (justified text)
        if para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            para_format.first_line_indent = Inches(0.5)
        
        # Set font for all runs
        for run in para.runs:
            run.font.name = 'Times New Roman'
            if not run.bold and not run.italic:
                run.font.size = Pt(12)

def format_headings(doc):
    """Format all headings consistently"""
    
    major_headings = [
        'INTRODUCTION', 'WHY BORS', 'NAME AND DESCRIPTION',
        'BUSINESS MODEL CANVAS', 'BROCHURE', 'MARKET VALIDATION',
        'FINANCIAL ASPECT', 'PROTOTYPING', 'PITCHING DIALOG',
        'CONCLUSION', 'APPENDICES', 'REFERENCES'
    ]
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Major section headings
        if text in major_headings or text.replace('?', '') in major_headings:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(24)
            para.paragraph_format.space_after = Pt(12)
            para.paragraph_format.keep_with_next = True
            
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(14)
                run.font.name = 'Times New Roman'
        
        # Sub-headings (like Value Proposition, Customer Segments, etc.)
        elif text in ['Value Proposition', 'Customer Segments', 'Channel', 
                      'Relationship', 'Key Partners', 'Cost', 'Key Resources',
                      'Manpower requirements', 'Organization structure',
                      'Key Activities', 'Revenue Stream', 'Projected Cost',
                      'Financial Statement – Project Cash Flow']:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.left_indent = Inches(0)
            
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'

def format_lists_and_bullets(doc):
    """Format lists with proper indentation"""
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # If paragraph starts with bullet or number
        if text.startswith('•') or text.startswith('-') or (len(text) > 0 and text[0].isdigit() and '. ' in text[:5]):
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.first_line_indent = Inches(-0.25)
            para.paragraph_format.space_after = Pt(3)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

def format_financial_tables(doc):
    """Format financial sections with proper alignment"""
    
    for para in doc.paragraphs:
        text = para.text
        
        # If paragraph contains financial data with dots
        if '........' in text or 'PHP' in text and '...' in text:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.space_after = Pt(3)
            
            for run in para.runs:
                run.font.name = 'Courier New'  # Monospace for alignment
                run.font.size = Pt(11)

def add_proper_page_breaks(doc):
    """Ensure major sections start on new pages"""
    
    major_sections = [
        'BUSINESS MODEL CANVAS', 'BROCHURE', 'MARKET VALIDATION',
        'FINANCIAL ASPECT', 'PROTOTYPING', 'PITCHING DIALOG',
        'CONCLUSION'
    ]
    
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() in major_sections:
            # Add page break before major sections (except first one)
            if i > 50:  # Skip early sections like title page and TOC
                para.paragraph_format.page_break_before = True

def format_image_placeholders(doc):
    """Format image placeholder text"""
    
    for para in doc.paragraphs:
        if '[INSERT' in para.text and 'IMAGE' in para.text:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(12)
            
            # Add border effect using paragraph shading
            para.paragraph_format.left_indent = Inches(1)
            para.paragraph_format.right_indent = Inches(1)
            
            for run in para.runs:
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(100, 100, 100)

def format_title_page(doc):
    """Format title page elements"""
    
    title_keywords = ['BORS', 'Barangay Online Records System', 
                      'A Project Portfolio', 'Presented to', 'Submitted to',
                      'In Partial Fulfillment', 'By']
    
    for para in doc.paragraphs[:40]:  # Only check first 40 paragraphs
        text = para.text.strip()
        
        # Center align title page elements
        if any(keyword in text for keyword in title_keywords):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(6)
        
        # Make title larger
        if 'BORS' in text and 'Barangay Online Records System' in text:
            for run in para.runs:
                run.font.size = Pt(18)
                run.bold = True

def format_table_of_contents(doc):
    """Format table of contents"""
    
    from docx.enum.text import WD_TAB_ALIGNMENT
    
    toc_found = False
    for i, para in enumerate(doc.paragraphs):
        if 'TABLE OF CONTENTS' in para.text:
            toc_found = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(18)
            
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(14)
            
            # Format next 30 paragraphs as TOC entries
            for j in range(i+1, min(i+31, len(doc.paragraphs))):
                toc_para = doc.paragraphs[j]
                if toc_para.text.strip():
                    toc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    toc_para.paragraph_format.space_after = Pt(3)
                    
                    # Add tab stops for page numbers
                    try:
                        tab_stops = toc_para.paragraph_format.tab_stops
                        tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT)
                    except:
                        pass  # Skip if tab stop fails
            
            break

def set_margins(doc):
    """Set proper margins for the document"""
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

def format_quotes_and_dialogs(doc):
    """Format quoted text and dialogs"""
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # If paragraph contains quotes
        if text.startswith('"') or text.startswith("'"):
            para.paragraph_format.left_indent = Inches(0.75)
            para.paragraph_format.right_indent = Inches(0.75)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            
            for run in para.runs:
                run.italic = True

print("Applying professional formatting to BORS proposal...")
print("Loading document...")

doc = Document('BORS_Technopreneurship_Proposal_FINAL.docx')

print("Setting margins...")
set_margins(doc)

print("Formatting title page...")
format_title_page(doc)

print("Formatting table of contents...")
format_table_of_contents(doc)

print("Formatting headings...")
format_headings(doc)

print("Applying paragraph formatting...")
format_document_professionally(doc)

print("Formatting lists and bullets...")
format_lists_and_bullets(doc)

print("Formatting financial sections...")
format_financial_tables(doc)

print("Formatting image placeholders...")
format_image_placeholders(doc)

print("Formatting quotes and dialogs...")
format_quotes_and_dialogs(doc)

print("Adding page breaks...")
add_proper_page_breaks(doc)

print("Saving formatted document...")
doc.save('BORS_Technopreneurship_Proposal_FORMATTED.docx')

print("\n✅ Professionally formatted proposal created!")
print("📄 File: BORS_Technopreneurship_Proposal_FORMATTED.docx")
print("\n📋 Formatting applied:")
print("   ✓ 1.5 line spacing throughout")
print("   ✓ Proper paragraph spacing (6pt after)")
print("   ✓ First-line indentation (0.5 inch)")
print("   ✓ Consistent heading styles")
print("   ✓ Professional margins (1.25\" left/right, 1\" top/bottom)")
print("   ✓ Times New Roman, 12pt body text")
print("   ✓ Centered title page elements")
print("   ✓ Formatted table of contents")
print("   ✓ Proper list indentation")
print("   ✓ Page breaks before major sections")
print("   ✓ Formatted image placeholders")
print("   ✓ Monospace font for financial tables")
print("\n✏️  Document is now publication-ready!")
