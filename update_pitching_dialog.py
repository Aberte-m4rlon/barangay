from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def update_pitching_dialog():
    """Update the Pitching Dialog section with BORS-specific content"""
    
    print("Updating Pitching Dialog section...")
    
    # Load the document
    doc = Document('BORS_Technopreneurship_Proposal_FINAL_WITH_IMAGES.docx')
    
    # Find the Pitching Dialog section
    pitching_found = False
    pitching_index = -1
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().upper()
        if 'PITCHING DIALOG' in text and not 'IMAGE' in text:
            pitching_found = True
            pitching_index = i
            print(f"✓ Found Pitching Dialog section at paragraph {i}")
            break
    
    if not pitching_found:
        print("❌ Pitching Dialog section not found")
        return
    
    # Find the next section to know where to stop
    next_section_index = -1
    for i in range(pitching_index + 1, len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip().upper()
        if text and (text.startswith('I.') or text.startswith('II.') or text.startswith('III.') or 
                     text.startswith('IV.') or text.startswith('V.') or text.startswith('VI.') or
                     text.startswith('VII.') or text.startswith('VIII.') or text.startswith('IX.') or
                     text.startswith('X.') or text.startswith('XI.') or text.startswith('XII.') or
                     'CONCLUSION' in text):
            next_section_index = i
            print(f"✓ Next section at paragraph {i}: {text[:50]}")
            break
    
    # Remove old content between pitching dialog and next section
    if next_section_index > 0:
        # Get elements to remove
        elements_to_remove = []
        for i in range(pitching_index + 1, next_section_index):
            para = doc.paragraphs[i]
            if para.text.strip():  # Only remove non-empty paragraphs
                elements_to_remove.append(para)
        
        # Remove from document
        for para in elements_to_remove:
            p = para._element
            p.getparent().remove(p)
        
        print(f"✓ Removed {len(elements_to_remove)} old paragraphs")
    
    # Insert new pitching dialog content after the heading
    pitching_heading = doc.paragraphs[pitching_index]
    
    # New pitching dialog content
    pitching_content = """According to the Philippine Statistics Authority, "Many barangays still rely on manual, paper-based systems for record-keeping, leading to inefficiencies, delays, and increased risk of data loss or corruption."

And these are some problems we encounter in traditional barangay governance:

• Do you find managing resident records overwhelming?
• Is it difficult for you to track blotter incidents efficiently?
• Do residents wait too long to get certificates?
• Do you spend countless hours on manual paperwork?
• Is communication between officials slow and ineffective?

But why would we suffer with outdated systems when we could make it easy!

And now… We're introducing to you BORS - Barangay Officials Record System!

✓ Without the hassle of searching through filing cabinets
✓ Without the risk of losing important documents
✓ Without consuming hours on manual data entry
✓ With just one click, you can access any resident record!
✓ With real-time updates and instant notifications!

This system can be used by all barangay officials but most especially for:
• Busy barangay captains managing multiple tasks
• Secretaries handling documentation
• Tanods tracking blotter records
• Treasurers managing financial records
• All officials needing secure communication

BORS Features at Your Fingertips:
📊 Real-time dashboard analytics
👥 Complete resident database management
📋 Digital blotter record system
📄 Automated certificate generation
📢 Community-wide announcements
💬 Secure officials chat with video calls
🔒 Role-based access control
📱 Mobile-responsive design

So what are you waiting for! Embrace digital transformation and modernize your barangay governance today!

"The future of governance is digital. Those who adapt will thrive, while those who resist will be left behind." - Digital Governance Initiative

Transform your barangay. Empower your community. Choose BORS."""

    # Insert the new content
    # We need to insert after the heading
    parent = pitching_heading._element.getparent()
    heading_index = parent.index(pitching_heading._element)
    
    # Split content into paragraphs
    lines = pitching_content.split('\n')
    
    for line_num, line in enumerate(lines):
        line = line.strip()
        if not line:
            # Add blank line
            new_para = doc.add_paragraph()
            new_para_element = new_para._element
            parent.insert(heading_index + 1 + line_num, new_para_element)
        else:
            # Add content line
            new_para = doc.add_paragraph()
            new_para.text = line
            
            # Style based on content
            if line.startswith('According to') or line.startswith('"'):
                # Quote style
                new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in new_para.runs:
                    run.font.italic = True
                    run.font.size = Pt(11)
            elif line.startswith('And these are') or line.startswith('But why') or line.startswith('And now'):
                # Emphasis style
                new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in new_para.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(46, 125, 50)
            elif line.startswith('•') or line.startswith('✓'):
                # Bullet points
                new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                new_para.paragraph_format.left_indent = Inches(0.5)
                for run in new_para.runs:
                    run.font.size = Pt(11)
            elif line.startswith('This system') or line.startswith('BORS Features') or line.startswith('So what are'):
                # Call to action
                new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in new_para.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
            elif line.startswith('Transform your'):
                # Final call to action
                new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in new_para.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(25, 118, 210)
            else:
                # Regular text
                new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in new_para.runs:
                    run.font.size = Pt(11)
            
            # Insert into document
            new_para_element = new_para._element
            parent.insert(heading_index + 1 + line_num, new_para_element)
    
    print(f"✓ Added new pitching dialog content ({len(lines)} lines)")
    
    # Save the document
    output_file = 'BORS_Technopreneurship_Proposal_COMPLETE_ALL_SECTIONS.docx'
    doc.save(output_file)
    print(f"\n✅ Document saved: {output_file}")
    print("\n📢 Pitching Dialog section updated with BORS-specific content!")
    print("\n🎯 New content includes:")
    print("   • Engaging opening quote")
    print("   • Problem identification")
    print("   • Solution introduction")
    print("   • Feature highlights")
    print("   • Target audience")
    print("   • Strong call-to-action")

if __name__ == "__main__":
    update_pitching_dialog()
